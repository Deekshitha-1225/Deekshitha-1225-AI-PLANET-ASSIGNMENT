from docx import Document

def create_proposal(company_name, industry_info, use_cases, datasets):
    doc = Document()
    
    # Company Introduction
    doc.add_heading(f"Proposal for {company_name}", 0)
    doc.add_paragraph(f"This proposal outlines relevant AI/ML use cases for {company_name} in the {industry_info[0]} industry.")
    
    # Use Cases
    doc.add_heading('Use Cases', level=1)
    for i, case in enumerate(use_cases, 1):
        doc.add_paragraph(f"{i}. {case}")
    
    # Resource Assets
    doc.add_heading('Relevant Datasets', level=1)
    for dataset in datasets:
        doc.add_paragraph(f"Title: {dataset['title']}")
        doc.add_paragraph(f"Link: {dataset['link']}")
    
    doc.save(f"{company_name}_AI_Proposal.docx")
    print(f"Proposal saved for {company_name}")

# Example
if __name__ == "__main__":
    company_name = "Tesla"
    industry_info = ["Automotive", "Electric Vehicles", "AI in Manufacturing"]
    use_cases = ["AI for predictive maintenance", "AI-powered customer experience"]
    datasets = [{"title": "EV Dataset", "link": "https://kaggle.com/dataset/ev"}]
    
    create_proposal(company_name, industry_info, use_cases, datasets)
